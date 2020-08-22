import docx
import os
from tkinter import *
from tkinter import ttk
from tkinter.filedialog import askopenfilename
from tkinter import filedialog
import tkinter as tk

#Get Taget File and Select Output Folder
Tk().withdraw()
targetFile = askopenfilename()
targetFolder = filedialog.askdirectory()


#Establish Lists
targetList = []
tscDocList = []
includedTSCList = []
exceptionCriteriaList = []
exceptionList = []
multipleExceptionsList = []
combinedExceptionsList = []
criteraDescriptionIndexList = []
exceptionIndexList = []
controlList = []
noExceptionsList = []
noGoList = []
tscDescList = []
exceptionTestList = []
otherTestList = []
otherTestIndexList = []
criteriaDescList = []

#Establish 2D Data List
dataList = [[], [], [], []]


'''
*******************************************
SECTION ONE: DOCUMENT SCAN AND LIST FILLING
*******************************************
'''


#Target Documents
tscDoc = docx.Document("TSC Sheet.docx")
tscDescDoc = docx.Document('TSC Desc List.docx')
targetDoc = docx.Document(targetFile)


#Turn documents into lists
for table in targetDoc.tables:
    for row in table.rows:
        for cell in row.cells:
            targetList.append(cell.text.strip())

for paragraph in tscDoc.paragraphs:
    tscDocList.append(paragraph.text)

for paragraph in tscDescDoc.paragraphs:
    tscDescList.append(paragraph.text)

#Remove Blank Spaces in Target Doc List
while("" in targetList) :
    targetList.remove("")

#Turn the tscDocList into a set for comparison to tar
tscSet = set(tscDocList)

#Make a List of All Included Criteria
for i in tscDocList:
    if i in targetList:
        includedTSCList.append(i)

'''
**************************************
SECTION 2: DATA COLLECTION
**************************************
'''

'''
PART 1: Get Exceptions and Indexes
'''

#Create Lists of Exceptions and Combine Them
exceptionList = [i for i in targetList if "Exception noted. " in i]
multipleExceptionsList = [i for i in targetList if "Exceptions noted. " in i]
combinedExceptionsList = exceptionList + multipleExceptionsList

#Get Indexes for Exceptions
for x, element in enumerate(targetList):
    if element in combinedExceptionsList:
        exceptionIndexList.append(x)

'''
PART 2: Get Testing and Indexes
'''

#Make List of All Indexes for Testing
exceptionTestingIndexList = list.copy(exceptionIndexList)

for i in range(len(exceptionTestingIndexList)):
    exceptionTestingIndexList[i] = exceptionTestingIndexList[i] - 1

for x in exceptionTestingIndexList:
        exceptionTestList.append(targetList[x])

#Make List of No Exception Phrases
noExceptionsList = ['No exceptions noted.', 'No exception noted.']

#Get all non-exception test indexes
for i, x in enumerate(targetList):
    if x in noExceptionsList:
        otherTestIndexList.append(i)

for i in range(len(otherTestIndexList)):
    otherTestIndexList[i] = otherTestIndexList[i] - 1

#Get All Non-Exception Tests
for i, x in enumerate(targetList):
    if i in otherTestIndexList:
        otherTestList.append(targetList[i])


'''
PART 3: Get Exception Criteria
'''

#Find All of the Excption Criteria
for exceptionIndex in exceptionIndexList:
    for targetIndex, notUsed in enumerate(targetList):
        if exceptionIndex == targetIndex:
            item = targetList[exceptionIndex]
            counter = 1
            notFound = True
            while notFound == True:
                if item not in includedTSCList:
                    item = targetList[exceptionIndex - counter]
                    counter += 1
                else:
                    exceptionCriteriaList.append(item)
                    notFound = False

'''
PART 4: Get Criteria Descriptions
'''

#Get Exception Criteria Desc. Indexes
for x, element in enumerate(targetList):
    for i in exceptionCriteriaList:
        if element == i:
          criteraDescriptionIndexList.append(x)

for i in range(len(criteraDescriptionIndexList)):
    criteraDescriptionIndexList[i] = criteraDescriptionIndexList[i] + 1

#Get Exception Criteria Descriptions
for x in criteraDescriptionIndexList:
        criteriaDescList.append(targetList[x])


'''
PART 5: Get All Exception Controls
'''

#Make List of Items for the Exception Control Search to Ignore
noGoList = noExceptionsList + combinedExceptionsList + includedTSCList + tscDescList + exceptionTestList + otherTestList

#Find all of the Exception Controls
for exceptionIndex in exceptionIndexList:
    for targetIndex, notUsed in enumerate(targetList):
        if exceptionIndex == targetIndex:
            item = targetList[exceptionIndex]
            counter = 1
            notFound = True
            while notFound == True:
                if item in noGoList:
                    item = targetList[exceptionIndex - counter]
                    counter += 1
                else:
                    controlList.append(item)
                    notFound = False

'''
**************************************
SECTION 3: AGGREGATE 
**************************************
'''

'''
PART 1: Data List
'''

#Put Criteria In Data List
for i in exceptionCriteriaList:
    dataList[0].append(i)

#Put Criteria Desc. in Data List
for i in criteriaDescList:
    dataList[1].append(i)

#Put Control in DataList
for i in controlList:
    dataList[2].append(i)

#Put Exceptions in Data List
for i in targetList:
    for x in exceptionIndexList:
        if x == targetList.index(i):
            dataList[3].append(i)

'''
Part 2: Non Duplicate Lists
'''

#Make Function to Remove Duplicates from list
def removeDuplicates(aList):
    output = list(dict.fromkeys(aList))
    return output

#Non Duplicate Criteria List
noDupExceptionCriteriaList = removeDuplicates(exceptionCriteriaList)

#Non Duplicate Description List
noDupCriteriaDescList = removeDuplicates(criteriaDescList)

#Non Duplicate Control List
noDupControlList = removeDuplicates(controlList)

#Non Duplicate Testing List
noDupExceptionList = removeDuplicates(dataList[3])

'''
**********************
SECTION 4: OUTPUT DATA
**********************
'''

#Create New Document Give Title
newDoc = docx.Document()
newDoc.add_paragraph('WP-13 Cheat Sheet')

#Give Criteria + Desc List
newDoc.add_paragraph('Affected Criteria')
for i in range(len(noDupCriteriaDescList)):
    newDoc.add_paragraph(noDupExceptionCriteriaList[i] + " - " + noDupCriteriaDescList[i])
newDoc.add_page_break()

#Give Affected Controls List
newDoc.add_paragraph('Affected Controls')
for i in range(len(noDupControlList)):
    newDoc.add_paragraph(noDupControlList[i])
newDoc.add_page_break()

#Give Results of Tests
newDoc.add_paragraph('Results of Tests')
for i in range(len(noDupExceptionList)):
    newDoc.add_paragraph(noDupExceptionList[i])
newDoc.add_page_break()

#Insert Data from Data List Into Document
newDoc.add_paragraph('Exception(s) With Criteria and Control')

for i in range(len(exceptionCriteriaList)):
    newDoc.add_paragraph(str(i+1) + ".")
    newDoc.add_paragraph(dataList[0][i] + " - " + dataList[1][i])
    newDoc.add_paragraph("Affected Control - " + dataList[2][i])
    newDoc.add_paragraph("Results of Test - " + dataList[3][i])
    newDoc.add_paragraph("")

#Save Document
newDoc.save(targetFolder + '/Cheat Sheet.docx')


